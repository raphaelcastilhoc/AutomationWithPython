import docx
from os import path
import Helpers.TextExtractor as TextExtractor

filePath = input('Type the file path: ')

if not path.isfile(filePath):
    print('%s is not a valid file path.' % filePath)
    quit()

wordFile = docx.Document(filePath)

contacts = []
for paragraph in wordFile.paragraphs:
    wordContacts = TextExtractor.extractContacts(paragraph.text)
    for wordContact in wordContacts:
        contacts.append(wordContact)

if not contacts:
    print('Contacts not found')
    quit()

newWordFile = docx.Document()
for contact in contacts:
    newWordFile.add_paragraph(contact)

fileName, fileExtension = path.splitext(filePath)
newWordFileName = '%s_only_contacts.docx' % fileName
newWordFile.save(newWordFileName)

print('Contacts extracted to %s' % newWordFileName)